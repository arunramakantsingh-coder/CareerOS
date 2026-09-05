from __future__ import annotations

import hashlib
import io
import re
import subprocess
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from app.utils.document_intelligence_v2 import classify_document

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".rtf", ".jpg", ".jpeg", ".png", ".tif", ".tiff", ".webp", ".zip"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".webp"}
MAX_FILE_SIZE = 25 * 1024 * 1024
MAX_ARCHIVE_SIZE = 100 * 1024 * 1024
MAX_ARCHIVE_UNCOMPRESSED_SIZE = 250 * 1024 * 1024
MAX_ARCHIVE_FILES = 100


def sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def safe_filename(name: str) -> str:
    name = Path(name or "document").name
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" .")
    return name[:255] or "document"


def safe_relative_path(relative_path: str | None) -> str | None:
    if not relative_path:
        return None
    normalized = relative_path.replace("\\", "/").lstrip("/")
    parts = [p for p in normalized.split("/") if p not in ("", ".")]
    if any(p == ".." for p in parts):
        raise ValueError("Unsafe relative path")
    return "/".join(safe_filename(p) for p in parts)


def _usable_text(text: str) -> bool:
    value = re.sub(r"\s+", " ", text or "").strip()
    if len(value) < 60:
        return False
    alpha = sum(ch.isalpha() for ch in value)
    return alpha >= max(30, int(len(value) * 0.30))


def _ocr_image(image) -> str:
    """Run a small set of OCR passes suitable for scanned career/identity documents."""
    import pytesseract
    from PIL import ImageEnhance, ImageFilter, ImageOps

    image = image.convert("L")
    # Upscaling is particularly important for small PAN/Aadhaar cards and compressed scans.
    scale = 3 if max(image.size) < 2200 else 2
    image = image.resize((image.width * scale, image.height * scale))
    image = ImageOps.autocontrast(image)
    variants = [image, image.point(lambda p: 255 if p > 170 else 0), ImageEnhance.Contrast(image).enhance(1.6).filter(ImageFilter.SHARPEN)]
    texts: list[str] = []
    for variant in variants:
        for config in ("--psm 6", "--psm 11"):
            try:
                value = pytesseract.image_to_string(variant, config=config).strip()
                if value and value not in texts:
                    texts.append(value)
            except Exception:
                continue
    return "\n\n".join(texts)


def _pdf_layout_text(content: bytes) -> tuple[str, int]:
    """Extract PDF text using PyMuPDF blocks while preserving page/column boundaries."""
    import fitz
    pdf = fitz.open(stream=content, filetype="pdf")
    pages: List[str] = []
    for page in pdf:
        blocks = [b for b in page.get_text("blocks") if len(b) >= 5 and str(b[4]).strip()]
        if not blocks:
            pages.append("")
            continue
        width = float(page.rect.width or 1)
        full = [b for b in blocks if (b[2] - b[0]) >= width * 0.72]
        left = [b for b in blocks if b not in full and b[0] < width * 0.52]
        right = [b for b in blocks if b not in full and b[0] >= width * 0.52]
        two_columns = len(left) >= 2 and len(right) >= 2 and (min(b[0] for b in right) - max(b[2] for b in left)) > width * 0.04
        if two_columns:
            # Keep each page's columns as explicit blocks. The parser can then reject sidebar
            # competency strips instead of treating them as part of an employment record.
            ordered: List[Any] = sorted(full, key=lambda b: (b[1], b[0])) + sorted(left, key=lambda b: (b[1], b[0])) + sorted(right, key=lambda b: (b[1], b[0]))
        else:
            ordered = sorted(blocks, key=lambda b: (b[1], b[0]))
        pages.append("\n".join(str(b[4]).strip() for b in ordered))
    return "\n\n".join(pages).strip(), len(pdf)


def _pdf_ocr(content: bytes) -> tuple[str, int]:
    import fitz
    from PIL import Image
    pdf = fitz.open(stream=content, filetype="pdf")
    pages: List[str] = []
    for page in pdf:
        pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        pages.append(_ocr_image(image))
    return "\n\n".join(pages).strip(), len(pdf)


def extract_text(filename: str, mime_type: str | None, content: bytes) -> Tuple[str, Dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    meta: Dict[str, Any] = {"method": "none", "ocr_required": False, "page_count": None}
    if suffix == ".txt" or (mime_type or "").startswith("text/"):
        return content.decode("utf-8", errors="replace"), {**meta, "method": "text"}
    if suffix == ".rtf":
        try:
            plain = re.sub(r"\\[a-z]+\d* ?|[{}]", " ", content.decode("utf-8", errors="replace"))
            return re.sub(r"\s+", " ", plain).strip(), {**meta, "method": "rtf_text"}
        except Exception as exc:
            return "", {**meta, "method": "rtf_error", "error": str(exc)}
    if suffix == ".doc":
        try:
            result = subprocess.run(["antiword", "-"], input=content, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=15, check=True)
            return result.stdout.decode("utf-8", errors="replace"), {**meta, "method": "antiword"}
        except Exception as exc:
            return "", {**meta, "method": "antiword_error", "error": str(exc)}
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Include table cells; many offer letters and resumes use tables for headers/dates.
            for table in doc.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        paragraphs.append(" | ".join(values))
            return "\n".join(paragraphs), {**meta, "method": "docx"}
        except Exception as exc:
            return "", {**meta, "method": "docx_error", "error": str(exc)}
    if suffix == ".pdf":
        pypdf_text = ""
        pypdf_error = None
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pypdf_text = "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
            meta["page_count"] = len(reader.pages)
        except Exception as exc:
            pypdf_error = str(exc)
        layout_text = ""
        try:
            layout_text, page_count = _pdf_layout_text(content)
            meta["page_count"] = page_count
        except Exception as exc:
            meta["layout_error"] = str(exc)

        candidate = layout_text if _usable_text(layout_text) else pypdf_text
        # Short text layers are frequently OCR remnants from scanned certificates/cards. Run OCR
        # when the extracted content is sparse; this is the key safeguard against filename-only UX.
        if _usable_text(candidate) and len(candidate) >= 500:
            return candidate, {**meta, "method": "pdf_layout" if candidate == layout_text else "pdf_text", "layout_aware": candidate == layout_text}
        try:
            ocr_text, page_count = _pdf_ocr(content)
            meta["page_count"] = page_count
            if len(ocr_text) > len(candidate) or not _usable_text(candidate):
                return ocr_text, {**meta, "ocr_required": True, "method": "pdf_ocr", "pypdf_error": pypdf_error, "fallback_from": "pdf_layout" if candidate == layout_text else "pdf_text"}
        except Exception as exc:
            meta["ocr_error"] = str(exc)
        if candidate:
            return candidate, {**meta, "method": "pdf_layout" if candidate == layout_text else "pdf_text", "layout_aware": candidate == layout_text, "ocr_required": False, "pypdf_error": pypdf_error}
        return pypdf_text, {**meta, "ocr_required": True, "method": "pdf_ocr_error", "pypdf_error": pypdf_error}
    if suffix in IMAGE_EXTENSIONS:
        try:
            from PIL import Image
            image = Image.open(io.BytesIO(content))
            text = _ocr_image(image)
            return text.strip(), {**meta, "method": "image_ocr", "ocr_required": True, "image_width": image.width, "image_height": image.height}
        except Exception as exc:
            return "", {**meta, "method": "image_ocr_error", "ocr_required": True, "error": str(exc)}
    return "", {**meta, "method": "unsupported"}


def image_to_pdf(content: bytes) -> bytes:
    from PIL import Image
    image = Image.open(io.BytesIO(content))
    if getattr(image, "is_animated", False):
        image.seek(0)
    if image.mode not in ("RGB", "L"):
        background = Image.new("RGB", image.size, "white")
        if "A" in image.getbands():
            background.paste(image, mask=image.getchannel("A"))
        else:
            background.paste(image.convert("RGB"))
        image = background
    elif image.mode == "L":
        image = image.convert("RGB")
    output = io.BytesIO()
    image.save(output, format="PDF", resolution=150.0)
    return output.getvalue()


def build_markdown_record(*, document_id: str, owner: str | None, original_filename: str, stored_filename: str, content_hash: str, classification: Dict[str, Any], extraction_meta: Dict[str, Any], extracted_text: str, relative_path: str | None, derived_pdf_path: str | None) -> str:
    category = classification.get("category", "other")
    subtype = classification.get("subcategory", "unclassified")
    confidence = classification.get("confidence", 0)
    lines = [f"# CareerOS Evidence Record — {original_filename}", "", "## Identity", f"- Document ID: `{document_id}`", f"- Owner: {owner or 'Unknown'}", f"- Original filename: `{original_filename}`", f"- Stored filename: `{stored_filename}`", f"- Source relative path: `{relative_path or original_filename}`", f"- SHA-256: `{content_hash}`", "", "## Classification", f"- Category: **{category}**", f"- Subtype: **{subtype}**", f"- Confidence: **{confidence}**", "", "## Extraction", f"- Method: `{extraction_meta.get('method', 'none')}`", f"- OCR required: `{bool(extraction_meta.get('ocr_required'))}`", f"- Page count: `{extraction_meta.get('page_count')}`"]
    if extraction_meta.get("layout_aware"):
        lines.append("- Layout aware: `true`")
    if extraction_meta.get("image_width") and extraction_meta.get("image_height"):
        lines.append(f"- Image dimensions: `{extraction_meta.get('image_width')} x {extraction_meta.get('image_height')}`")
    if derived_pdf_path:
        lines.extend(["", "## Derived Artifacts", f"- Normalized PDF: `{derived_pdf_path}`", "- Original evidence remains authoritative."])
    lines.extend(["", "## Extracted Text", "", extracted_text[:100000] if extracted_text else "_No extractable text was produced._", "", "## Provenance", "This Markdown file is a derived CareerOS index/audit artifact. The original uploaded document remains the authoritative evidence.", ""])
    return "\n".join(lines)


def canonical_filename(owner: str | None, classification: Dict[str, Any], issuer: str | None, original: str) -> str:
    stem = Path(original).stem
    owner_part = re.sub(r"[^A-Za-z0-9 -]+", "", owner or "Professional").strip()
    category = classification.get("category", "document").replace("_", " ").title()
    subtype = classification.get("subcategory", "").replace("_", " ").title()
    issuer_part = re.sub(r"[^A-Za-z0-9 -]+", "", issuer or "").strip()
    pieces = [owner_part, subtype or category]
    if issuer_part:
        pieces.append(issuer_part)
    if stem and stem.lower() not in {"document", "scan", "img", "image"}:
        pieces.append(stem)
    suffix = Path(original).suffix.lower() or ".bin"
    base = safe_filename(" - ".join(p for p in pieces if p))[:230]
    return f"{base}{suffix}"


def iter_zip_entries(content: bytes) -> Iterable[Tuple[str, bytes]]:
    if len(content) > MAX_ARCHIVE_SIZE:
        raise ValueError("ZIP archive exceeds the 100MB limit")
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        infos = [info for info in archive.infolist() if not info.is_dir()]
        if len(infos) > MAX_ARCHIVE_FILES:
            raise ValueError("ZIP archive exceeds the 100-file limit")
        total_uncompressed = sum(info.file_size for info in infos)
        if total_uncompressed > MAX_ARCHIVE_UNCOMPRESSED_SIZE:
            raise ValueError("ZIP archive expands beyond the 250MB safety limit")
        for info in infos:
            relative = safe_relative_path(info.filename) or "document"
            suffix = Path(relative).suffix.lower()
            if suffix not in SUPPORTED_EXTENSIONS or suffix == ".zip":
                continue
            if info.file_size > MAX_FILE_SIZE:
                continue
            if any(part == ".." for part in Path(relative).parts):
                raise ValueError("Unsafe ZIP entry path")
            yield relative, archive.read(info)
