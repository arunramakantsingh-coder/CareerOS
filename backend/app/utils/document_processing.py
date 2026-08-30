from __future__ import annotations

import hashlib
import io
import os
import re
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader
import pytesseract

STORAGE_ROOT = Path(os.getenv('CAREEROS_STORAGE_ROOT', '/app/storage')).resolve()
MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_ARCHIVE_BYTES = 100 * 1024 * 1024
MAX_ARCHIVE_FILES = 100
SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt', '.md', '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.xlsx', '.xls', '.zip'}


def safe_filename(name: str) -> str:
    name = name.replace('\\', '/').split('/')[-1]
    name = re.sub(r'[^A-Za-z0-9._()\- ]+', '_', name).strip(' .')
    return name[:180] or 'document'


def safe_archive_name(name: str) -> str:
    normalized = name.replace('\\', '/')
    if normalized.startswith('/') or any(part in {'..', ''} and part == '..' for part in normalized.split('/')):
        raise ValueError('Unsafe archive path')
    return normalized


def sha256_bytes(content: bytes) -> str: return hashlib.sha256(content).hexdigest()
def is_supported(name: str) -> bool: return Path(name).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_zip(content: bytes) -> list[tuple[str, bytes]]:
    if len(content) > MAX_ARCHIVE_BYTES: raise ValueError('Archive exceeds the maximum allowed size')
    result: list[tuple[str, bytes]] = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        infos = [info for info in archive.infolist() if not info.is_dir()]
        if len(infos) > MAX_ARCHIVE_FILES: raise ValueError('Archive contains too many files')
        for info in infos:
            name = safe_archive_name(info.filename)
            if not is_supported(name): continue
            if info.file_size > MAX_FILE_BYTES: raise ValueError(f'Archive member exceeds the file size limit: {name}')
            result.append((name, archive.read(info)))
    return result


def save_document(candidate_id: str, original_name: str, content: bytes) -> tuple[str, str]:
    if len(content) > MAX_FILE_BYTES: raise ValueError('File exceeds the maximum allowed size of 25 MB')
    folder = STORAGE_ROOT / 'documents' / str(candidate_id); folder.mkdir(parents=True, exist_ok=True)
    digest = sha256_bytes(content); filename = f'{digest[:16]}-{safe_filename(original_name)}'; path = folder / filename
    path.write_bytes(content); return str(path.relative_to(STORAGE_ROOT)), digest


def _ocr_pdf(reader: PdfReader) -> str:
    chunks: list[str] = []
    for page in reader.pages[:25]:
        for image_file in page.images:
            try:
                image = image_file.image or Image.open(io.BytesIO(image_file.data))
                chunks.append(pytesseract.image_to_string(image))
            except Exception:
                continue
    return '\n'.join(chunks)


def read_document_text(path: str) -> tuple[str, str]:
    full_path = (STORAGE_ROOT / path).resolve()
    if STORAGE_ROOT not in full_path.parents or not full_path.is_file(): raise FileNotFoundError('Stored document is unavailable')
    suffix = full_path.suffix.lower()
    if suffix in {'.txt', '.md'}: return full_path.read_text(encoding='utf-8', errors='ignore'), 'text'
    if suffix == '.pdf':
        reader = PdfReader(str(full_path)); text = '\n'.join((page.extract_text() or '') for page in reader.pages)
        if text.strip(): return text, 'pdf-text'
        ocr = _ocr_pdf(reader)
        return ocr, 'pdf-ocr' if ocr.strip() else 'pdf-no-text'
    if suffix == '.docx':
        doc = DocxDocument(str(full_path)); parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows: parts.append(' | '.join(cell.text for cell in row.cells))
        return '\n'.join(parts), 'docx'
    if suffix in {'.jpg', '.jpeg', '.png', '.tif', '.tiff'}:
        return pytesseract.image_to_string(Image.open(full_path)), 'ocr'
    return '', 'unsupported-text'
